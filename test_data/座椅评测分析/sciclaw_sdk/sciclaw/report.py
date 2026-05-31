"""
SciClaw 专业报告生成器
======================
生成含嵌入式图表的 Word(.docx) 和 Markdown 报告。
"""

import os, time, json
from datetime import datetime

def generate(result, output_path: str, format: str = "docx",
             include_charts: bool = True, charts_dir: str = None) -> str:
    """
    生成综合分析报告。

    Args:
        result: AnalysisResult 对象
        output_path: 输出路径
        format: "docx" | "md"
        include_charts: 是否嵌入图表
        charts_dir: 图表目录

    Returns:
        输出文件路径
    """
    if format == "md":
        return _generate_markdown(result, output_path, include_charts, charts_dir)
    else:
        return _generate_docx(result, output_path, include_charts, charts_dir)


def _generate_markdown(result, path, include_charts, charts_dir):
    """生成Markdown报告"""
    lines = []
    info = result.info

    lines.append(f"# 乘用车座椅综合评测 — 分析报告")
    lines.append(f"\n**生成时间**: {result.timestamp}")
    lines.append(f"**数据时长**: {info.get('duration_s','?')}s")
    lines.append(f"**车速**: {info.get('speed_range','?')}")
    lines.append(f"**方向盘**: {info.get('wheel_range','?')}")

    # 事件
    lines.append(f"\n## 1. 驾驶行为事件 ({len(result.events)}个)")
    if result.events:
        lines.append(f"\n| 事件类型 | 时间 | 持续 |")
        lines.append(f"|---------|------|------|")
        for e in result.events:
            name = getattr(e, 'event_name', str(e))
            t_s = getattr(e, 't_start', 0)
            t_e = getattr(e, 't_end', 0)
            lines.append(f"| {name} | {t_s:.1f}~{t_e:.1f}s | {t_e-t_s:.2f}s |")

    # 对比
    if result.comparison:
        lines.append(f"\n## 2. 实验组 vs 对照组")
        lines.append(f"\n| 指标 | 实验组 | 对照组 | 衰减率 | 结论 |")
        lines.append(f"|------|--------|--------|--------|------|")
        for name, vals in result.comparison.items():
            lines.append(f"| {name} | {vals.get('exp','?')} | {vals.get('ctrl','?')} | "
                        f"{vals.get('atten_pct',0):+.1f}% | {vals.get('verdict','')} |")

    # 图表
    if include_charts and charts_dir:
        lines.append(f"\n## 3. 分析图表")
        for name, chart_path in (result.charts or {}).items():
            rel = os.path.relpath(chart_path, os.path.dirname(path))
            lines.append(f"\n### {name}")
            lines.append(f"![{name}]({rel})")

    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    return path


def _generate_docx(result, path, include_charts, charts_dir):
    """生成Word报告 (docx)"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # 回退到Markdown
        print("python-docx 未安装，使用 Markdown 格式")
        return _generate_markdown(result, path.replace('.docx','.md'),
                                  include_charts, charts_dir)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    info = result.info

    # 标题
    title = doc.add_heading('乘用车座椅综合评测 — 分析报告', level=0)
    doc.add_paragraph(f"生成时间: {result.timestamp} | "
                      f"数据时长: {info.get('duration_s','?')}s | "
                      f"车速: {info.get('speed_range','?')}")

    # 第1章: 驾驶事件
    doc.add_heading('1. 驾驶行为事件', level=1)
    if result.events:
        table = doc.add_table(rows=len(result.events)+1, cols=4)
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(['事件类型','开始(s)','结束(s)','持续(s)']):
            table.rows[0].cells[i].text = h
        for j, e in enumerate(result.events):
            table.rows[j+1].cells[0].text = getattr(e, 'event_name', str(e))
            table.rows[j+1].cells[1].text = f"{getattr(e,'t_start',0):.1f}"
            table.rows[j+1].cells[2].text = f"{getattr(e,'t_end',0):.1f}"
            table.rows[j+1].cells[3].text = f"{getattr(e,'t_end',0)-getattr(e,'t_start',0):.2f}"

    # 第2章: 对比
    if result.comparison:
        doc.add_heading('2. 实验组 vs 对照组', level=1)
        n = len(result.comparison)
        table = doc.add_table(rows=n+1, cols=5)
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(['指标','实验组','对照组','衰减率','结论']):
            table.rows[0].cells[i].text = h
        for j, (name, vals) in enumerate(result.comparison.items()):
            table.rows[j+1].cells[0].text = name
            table.rows[j+1].cells[1].text = str(vals.get('exp','?'))
            table.rows[j+1].cells[2].text = str(vals.get('ctrl','?'))
            table.rows[j+1].cells[3].text = f"{vals.get('atten_pct',0):+.1f}%"
            table.rows[j+1].cells[4].text = vals.get('verdict','')

    # 第3章: 图表
    if include_charts and result.charts:
        doc.add_heading('3. 分析图表', level=1)
        for name, chart_path in result.charts.items():
            if os.path.exists(chart_path):
                doc.add_heading(name.replace('_',' '), level=2)
                try:
                    doc.add_picture(chart_path, width=Inches(5.5))
                except:
                    doc.add_paragraph(f"[图表: {chart_path}]")

    doc.save(path)
    return path
